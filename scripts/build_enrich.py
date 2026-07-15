#!/usr/bin/env python3
"""Enrich Mission Control data with evidence, costed levers, Ofsted/drivers per school,
KPI frameworks and document digests. Re-runnable: reads the base mission_data.json
(created by build_data.py) and overwrites the enrichment keys + the AI brief."""
import json, os

ROOT = '/sessions/kind-laughing-tesla/mnt/DfE'
NEET = os.path.join(ROOT, 'Milburn/school-profile copy')
OUT = '/sessions/kind-laughing-tesla/mnt/outputs/mission-control'

md = json.load(open(os.path.join(OUT, 'public/data/mission_data.json')))
ofsted = json.load(open(os.path.join(NEET, 'public/ofsted.json')))
drivers = json.load(open(os.path.join(NEET, 'public/school_drivers.json')))
neet_brief = json.load(open(os.path.join(NEET, 'netlify/functions/neet-brief.json')))
wwc = json.load(open(os.path.join(NEET, 'public/wwc_data.json')))
evid = json.load(open(os.path.join(NEET, 'public/evidence_data.json')))

def r1(x):
    return None if x is None else round(float(x), 1)

# ---- 1. Join Ofsted sub-judgements + drivers onto every school ----
GRADE = {1: 'Outstanding', 2: 'Good', 3: 'Requires improvement', 4: 'Inadequate'}
for area in md['areas'].values():
    for phase in ('secondary', 'primary'):
        for s in area['schools'][phase]:
            urn = str(s['urn'])
            o = ofsted.get(urn)
            if o:
                s['ofsted_sub'] = {k: o.get(k) for k in ('oe', 'lm', 'qe', 'pd', 'ba') if o.get(k)}
                s['ofsted_date'] = o.get('id')
            d = drivers.get(urn)
            if d:
                s['drivers'] = {k: r1(d.get(k)) for k in ('pa', 'abs', 'susp', 'perm')}

# ---- 2. Costed policy levers with modelled impact (Ready to Work modelling) ----
md['levers'] = neet_brief['policy_levers_with_modelled_impact']
md['leversNote'] = ('Modelled maximum reduction in the national 16-24 NEET count at full rollout, '
                    'with indicative national annual cost (£m). Ready to Work sprint modelling, July 2026. '
                    'Mission-area pilots would be a small fraction of the national cost.')

# ---- 3. Evidence library ----
md['evidence'] = {
    'source': evid.get('source_note'),
    'riskFactors': evid['risk_factors'],  # {'17-19': [{f, ppt}], '20-24': [...]}
    'phaseCards': [
        {'phase': p['title'], 'ages': p['ages'], 'blurb': p['blurb'],
         'cards': [{'stat': c.get('stat'), 'label': c.get('label'), 'source': c.get('source')} for c in p.get('cards', [])]}
        for p in evid.get('phases', [])
    ],
    'wwc': {
        'source': wwc.get('source'), 'year': wwc.get('year'),
        'headline': wwc.get('headline'),
        'attainmentGap': wwc.get('attainment_gap'),
        'engagement': wwc.get('engagement'),
        'recommendations': wwc.get('recommendations'),
        'lines': wwc.get('lines'),
    },
}

# ---- 4. KPI frameworks (Strand 2 doc + Mission Metrics framework) ----
import docx
sd = docx.Document(os.path.join(ROOT, 'Missions/Strategic & Planning/Missions Strand 2 Beyond the School Gates.docx'))
tables = sd.tables
def rows(t):
    return [[c.text.strip() for c in r.cells] for r in t.rows]

outcome = [{'kpi': r[0], 'baseline': r[1], 'y1': r[2], 'y3': r[3]} for r in rows(tables[2])[1:]]
collective = [{'kpi': r[0], 'measurement': r[1], 'y3': r[2]} for r in rows(tables[3])[1:]]
sustain = [{'kpi': r[0], 'measurement': r[1], 'y3': r[2]} for r in rows(tables[4])[1:]]
risks = [{'risk': r[0], 'rating': r[1], 'impact': r[2], 'mitigation': r[3]} for r in rows(tables[5])[1:]]
partners = [{'partner': r[0], 'role': r[1], 'ne': r[2], 'coastal': r[3]} for r in rows(tables[1])[1:]]

mmd = docx.Document(os.path.join(ROOT, 'Missions/Strategic & Planning/Mission Metrics (1).docx'))
framework = []
category = None
for r in rows(mmd.tables[0])[1:]:
    if len(set(r)) == 1:  # category header row (all cells identical)
        category = r[0]
        continue
    framework.append({'metric': r[0], 'what': r[1], 'source': r[2], 'dimension': r[3], 'frequency': r[4], 'category': category})

md['kpis'] = {
    'outcome': outcome, 'collective': collective, 'sustainability': sustain,
    'framework': framework,
    'note': 'Outcome / collective-impact / sustainability KPIs from the Strand 2 collective impact strategy (May 2026); metrics framework from the Mission Metrics doc.',
}
md['strand2Risks'] = risks
md['partners'] = partners

# ---- 5. Enrichment benchmarks + doc digests ----
ed = docx.Document(os.path.join(ROOT, 'Enrichment/School Enrichment Benchmarks - clean version 17 September.docx'))
benchmarks = []
for t in ed.tables[:8]:
    rr = rows(t)
    title = rr[0][0]
    summary = next((c for r in rr[1:] for c in r if c.startswith('Summary')), '')
    benchmarks.append({'benchmark': title, 'summary': summary.replace('Summary', '', 1).strip()[:400]})
md['enrichmentBenchmarks'] = benchmarks

md['docDigests'] = {
    'whitePaper': ('Every Child Achieving and Thriving (Feb 2026): ten-year plan; targets national average A8 of 50 by end of decade '
                   'and halving the disadvantage gap within a generation. Three shifts: narrow→broad (knowledge-rich curriculum + guaranteed '
                   'enrichment), plus inclusion/SEND (£3.7bn specialist capital, disadvantage funding moving beyond FSM) and phased delivery '
                   '2025-28 (align to best practice → prepare for reform → deliver). Enrichment Framework + Ofsted benchmarks from autumn; '
                   'new KS3 alliance for primary-secondary transition.'),
    'inclusionFramework': ('National Inclusion Framework (internal direction paper): minimum National Inclusion Standards for every school by 2028, '
                           '£15m evidence base via independent expert panel. Framed as a floor not a ceiling; key design risk is compliance-driven '
                           'floor standards. Mission Inclusion Labs generate the delivery evidence this framework needs.'),
    'enrichmentBenchmarksNote': 'Eight School Enrichment Benchmarks (Sept 2025) — the national framework the mission enrichment entitlement aligns to.',
}

json.dump(md, open(os.path.join(OUT, 'public/data/mission_data.json'), 'w'), indent=1)
print('mission_data.json', os.path.getsize(os.path.join(OUT, 'public/data/mission_data.json')))

# ---- 6. Rebuild the AI brief with enrichment ----
brief_path = os.path.join(OUT, 'netlify/functions/mission-brief.json')
brief = json.load(open(brief_path))
brief['policy_levers_with_modelled_impact'] = md['levers']
brief['levers_note'] = md['leversNote']
brief['evidence_library'] = {
    'neet_risk_factors_ppt': evid['risk_factors'],
    'key_stats': [c for p in md['evidence']['phaseCards'] for c in p['cards']][:18],
    'white_working_class_inquiry': {'headline': wwc.get('headline'), 'recommendations': wwc.get('recommendations')},
}
brief['kpi_framework'] = {'outcome': outcome, 'collective': collective, 'sustainability': sustain}
brief['strand2_partners'] = partners
brief['strand2_risks'] = risks
brief['enrichment_benchmarks'] = benchmarks
brief['policy_context'] = md['docDigests']
# per-school drivers for mission secondaries
for k, a in md['areas'].items():
    secs = a['schools']['secondary']
    for i, s in enumerate(brief['areas'][k]['secondary_schools']):
        src = next((x for x in secs if x['name'] == s['name']), None)
        if src:
            s['ofsted_overall'] = GRADE.get((src.get('ofsted_sub') or {}).get('oe'))
            s['drivers'] = src.get('drivers')
json.dump(brief, open(brief_path, 'w'), separators=(',', ':'))
print('brief bytes:', os.path.getsize(brief_path))
print('kpis:', len(outcome), len(collective), len(sustain), 'framework:', len(framework), 'benchmarks:', len(benchmarks))
